from collections import defaultdict
import pickle
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from .onto.leavedonto import OntoManager, LeavedOnto

new = 'New Words'
legend = 'Legend'
shared = 'Shared Words'
cur = 'Current'
prev = 'Previous'
absent = 'Not in Current'
total_word_count = 'Total Word Count: {}'
vocab = 'Vocabulary'
FIELDS = ['word', 'origin']


def get_selected_fields(om, entry):
    filtered = []
    for f in FIELDS:
        value = om.onto1.get_field_value(entry, f)
        value = str(value)
        filtered.append(value)
    return filtered


def gen_vocab_report(onto_path, out_path):
    total_data = gather_total_data(onto_path)
    lessons_data = gather_lesson_data(onto_path)
    level = onto_path.stem

    # format it in docx
    total_file = out_path / f'{level} Vocab Report - Total.docx'
    export_total_vocab_report(level, total_data, total_file)

    lessons_file = out_path / f'{level} Vocab Report - Lessons.docx'
    export_lessons_vocab_report(level, lessons_data, lessons_file)


def export_lessons_vocab_report(level, total_data, out_file):
    doc = Document()
    styles = doc.styles

    freq_style = styles.add_style('freq', WD_STYLE_TYPE.CHARACTER)
    freq_font = freq_style.font
    freq_font.name = 'Lato'
    freq_font.italic = True

    entry_style = styles.add_style('entry', WD_STYLE_TYPE.CHARACTER)
    entry_font = entry_style.font
    entry_font.size = Pt(7)
    entry_font.name = 'Lato Light'

    tree_style = styles.add_style('tree', WD_STYLE_TYPE.CHARACTER)
    tree_font = tree_style.font
    tree_font.name = 'Lato Light'
    tree_font.size = Pt(11)

    # LESSONS
    for section, data in total_data.items():
        doc.add_heading(section, 0)
        for subtitle, subdata in data.items():
            if subtitle == absent:
                doc.add_heading(subtitle.replace('Current', section), 1)
            else:
                doc.add_heading(subtitle, 1)

            if subtitle.startswith('Total'):
                par = doc.add_paragraph()
                for level in subdata:
                    line = '\t'.join(level)
                    run = par.add_run(line, style=tree_style)
                    if ':' in line:
                        run.bold = True
                    run.add_break()
            elif subtitle == new or subtitle == absent:
                for ssubtitle, ssubdata in subdata.items():
                    doc.add_heading(ssubtitle, 3)
                    par = doc.add_paragraph()
                    for num, entry in enumerate(ssubdata):
                        run = par.add_run(f'{num+1}. ')
                        run.font.size = Pt(9)
                        par.add_run(' '.join(entry), style=entry_style)
                        par.add_run(' ')
            elif subtitle == shared:
                for ssubtitle, ssubdata in subdata.items():
                    doc.add_heading(ssubtitle, 3)
                    par = doc.add_paragraph()
                    for sssubtitle, sssubdata in ssubdata.items():
                        run = par.add_run(sssubtitle, style=entry_style)
                        run.bold = True
                        par.add_run()
                        for num, entry in enumerate(sssubdata):
                            run = par.add_run(f'{num+1}. ')
                            run.font.size = Pt(9)
                            par.add_run(' '.join(entry), style=entry_style)
                            par.add_run(' ')
                        par.add_run().add_break()
                    par.runs[-1].text = ''
            else:
                raise ValueError('this is unexpected!')

    doc.save(out_file)


def export_total_vocab_report(level, total_data, out_file):
    doc = Document()
    styles = doc.styles

    freq_style = styles.add_style('freq', WD_STYLE_TYPE.CHARACTER)
    freq_font = freq_style.font
    freq_font.name = 'Lato'
    freq_font.italic = True

    entry_style = styles.add_style('entry', WD_STYLE_TYPE.CHARACTER)
    entry_font = entry_style.font
    entry_font.size = Pt(7)
    entry_font.name = 'Lato Light'

    tree_style = styles.add_style('tree', WD_STYLE_TYPE.CHARACTER)
    tree_font = tree_style.font
    tree_font.name = 'Lato Light'
    tree_font.size = Pt(11)

    # TITLE
    doc.add_heading(f'{level} Vocabulary Report - Total', 0)

    # TOTAL
    for section, data in total_data.items():
        doc.add_heading(section, 1)
        if section.startswith('Total'):
            par = doc.add_paragraph()
            for level in data:
                line = '\t'.join(level)
                run = par.add_run(line, style=tree_style)
                if ':' in line:
                    run.bold = True
                run.add_break()

        else:
            for subtitle, subdata in data.items():
                doc.add_heading(subtitle, 2)
                par = doc.add_paragraph()
                for num, entry in enumerate(subdata):
                    run = par.add_run(f'{num+1}. ')
                    run.font.size = Pt(9)
                    par.add_run(' '.join(entry), style=entry_style)
                    par.add_run(' ')

    doc.save(out_file)


def gather_total_data(onto_path):
    om = OntoManager()
    for f in sorted(list(onto_path.glob('*.yaml'))):
        om.merge_to_onto(f)

    total_data = {}

    tree, total_words = om.onto1.export_tree_report()
    total_data[total_word_count.format(total_words)] = tree

    # gather data
    entries = om.onto1.ont.find_entries()

    # organize it for report
    total_data[vocab] = {}
    for p, e in entries:
        filtered = []
        for entry in e:
            filtered.append(get_selected_fields(om, entry))
        total_data[vocab]['/'.join(p)] = filtered

    return total_data


def gather_lesson_data(onto_path):
    # group ontos belonging to the same lesson
    levels = defaultdict(list)
    for f in sorted(list(onto_path.glob('*.yaml'))):
        lesson, _ = f.stem.split('-')
        levels[lesson].append(f)

    # generate one onto for each lesson
    all_ontos = []
    for l, ontos in levels.items():
        om = None
        for onto in ontos:
            ont = LeavedOnto(onto)
            if not om:
                om = OntoManager()
                om.onto1.ont_path = f.parent / (l + '_')  # required to merge these ontos

            om.merge_to_onto(ont, in_to_organize=False)
        all_ontos.append((l, om))

    previous_onto = OntoManager()
    report_data = {}
    for num, ao in enumerate(all_ontos):
        lesson, om = ao
        if lesson not in report_data:
            report_data[lesson] = {}

        # VOCAB TREE
        tree, total_words = om.onto1.export_tree_report()
        report_data[lesson][total_word_count.format(total_words)] = tree

        # WORD LISTS
        # everything is new vocab
        if num == 0:
            # gather data
            entries = om.onto1.ont.find_entries()

            # organize it for report
            report_data[lesson][new] = {}
            for p, e in entries:
                filtered_fields = []
                for entry in e:
                    filtered_fields.append(get_selected_fields(om, entry))
                report_data[lesson][new]['/'.join(p)] = filtered_fields

        # split words of lesson in: New, Shared, Unseen
        else:
            # gather data
            previous_onto.merge_to_onto(all_ontos[num-1][1].onto1, add_origin=False)
            current_only, common, previous_only = om.diff_ontos(previous_onto.onto1)

            # organize it for report
            report_data[lesson][new] = {}
            for p, e in current_only:
                title = '/'.join(p)
                if title not in report_data[lesson][new]:
                    report_data[lesson][new][title] = []

                report_data[lesson][new][title].append(get_selected_fields(om, e))

            report_data[lesson][shared] = {}
            for c, p in common:
                title = '/'.join(c[0])
                if title not in report_data[lesson][shared]:
                    report_data[lesson][shared][title] = {cur: [], prev: []}

                report_data[lesson][shared][title][cur].append(get_selected_fields(om, c[1]))
                report_data[lesson][shared][title][prev].append(get_selected_fields(om, p[1]))

            report_data[lesson][absent] = {}
            for p, e in previous_only:
                title = '/'.join(p)
                if title not in report_data[lesson][absent]:
                    report_data[lesson][absent][title] = []

                report_data[lesson][absent][title].append(get_selected_fields(om, e))

    return report_data

